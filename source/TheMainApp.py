import tkinter as tk
from tkinter import messagebox, simpledialog, ttk
from docx import Document
from reportlab.pdfgen import canvas
import uuid
import requests
from bs4 import BeautifulSoup
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
import mimetypes
import os
import tkinter as tk
from tkinter import messagebox, simpledialog, ttk
from docx import Document
from reportlab.pdfgen import canvas
import uuid
from tkinter import *
import requests
import re as regex
from bs4 import BeautifulSoup
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from email.message import EmailMessage
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email.mime.image import MIMEImage
from email.mime.audio import MIMEAudio
import google.auth
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
import requests
import os
import base64
import mimetypes
import google.auth
from google.auth.transport.requests import Request
from google.oauth2 import service_account  # Adicione esta linha
from email.message import EmailMessage
import logging

EMPRESAS = {
    "12345678000195": "Empresa A",
    "98765432000196": "Empresa B",
}

def obter_data_atual():
    raw_req = requests.get('https://www.worldtimeserver.com/current_time_in_UTC.aspx')
    soupParser = BeautifulSoup(raw_req.content, 'lxml')
    date_time = soupParser.find('div', {'class': 'local-time'})
    date_today = date_time.find('h4').get_text().strip()
    return date_today

def replace_placeholders(doc, replacements):
    for paragraph in doc.paragraphs:
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in replacements.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)
    return doc

def create_contract(template_path, replacements):
    doc = Document(template_path)
    doc = replace_placeholders(doc, replacements)
    contract_id = uuid.uuid4()
    new_file_name = f'Contrato_{contract_id}.docx'
    doc.save(new_file_name)
    return new_file_name

def enviar_email_com_anexo(sender_email, sender_password, recipient_email, subject, body_text, attachment_filename):
    message = MIMEMultipart()
    message['From'] = sender_email
    message['To'] = recipient_email
    message['Subject'] = subject
    message.attach(MIMEText(body_text, 'plain'))

    with open(attachment_filename, "rb") as attachment:
        part = MIMEBase("application", "vnd.openxmlformats-officedocument.wordprocessingml.document")
        part.set_payload(attachment.read())
        encoders.encode_base64(part)
        part.add_header(
            "Content-Disposition",
            f"attachment; filename={os.path.basename(attachment_filename)}"
        )
        message.attach(part)

    try:
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
            server.login(sender_email, sender_password)
            server.sendmail(sender_email, recipient_email, message.as_string())
        print("E-mail enviado com sucesso!")
    except Exception as e:
        print(f"Erro ao enviar e-mail: {e}")

# PLaceholders
def gerar_documento(cnpj, razao_social, administrador, identidade, cpf, modulos, valor_parcela, numero_parcelas, email):
    data_atual = obter_data_atual()
    substituicoes = {
        "POSITION": razao_social,
        "N2": cnpj,
        "D3": administrador,
        "C4": identidade,
        "E5": cpf,
        "ON6": numero_parcelas,
        "PO7": valor_parcela,
        "R47": modulos,
        "ADX500": data_atual,
    }
    
    # the Word .docx file that are your template file with the placeholders words. 
    template_path = r"/path-to-template-contract.docx"
    new_file_name = create_contract(template_path, substituicoes)
    # edit your email body here
    enviar_email_com_anexo(
        sender_email="emailqueestaenviando@gmail.com", # sender email
        sender_password="a-sua-senha-de-aplicativo", # app password
        recipient_email=email,
        subject="Send by a robot: 🤖",
        body_text="Hello You! Lego Napoleon movie Written in noble gas-filled glass tubes Underlined in sparks I'll admit it's elaborate for a wakin' thought",
        attachment_filename=new_file_name
    )

    messagebox.showinfo("Sucesso", f"Documentos gerados e enviados para: {email}")
    messagebox.showinfo("Sucesso", f"Documentos gerados: {new_file_name}")

def validar_senha(senha):
    return senha == "senha123"  # the app password, you can edit it

def login():
    senha = simpledialog.askstring("Senha", "Digite a senha:", show='*')
    if validar_senha(senha):
        inicializar_interface()
    else:
        messagebox.showerror("Erro", "Senha incorreta!")
        root.destroy()

def buscar_razao_social():
    cnpj = entry_cnpj.get()
    razao_social = EMPRESAS.get(cnpj, "CNPJ não encontrado")
    entry_razao_social.delete(0, tk.END)
    entry_razao_social.insert(0, razao_social)

def inicializar_interface():
    global entry_cnpj, entry_razao_social, entry_administrador, entry_identidade, entry_cpf, entry_valor_parcela, entry_numero_parcelas, modulos_var
    
    app = tk.Tk()
    app.title("Gerador de Contratos")

    style = ttk.Style()
    style.configure("TLabel", font=("Helvetica", 10))
    style.configure("TEntry", font=("Helvetica", 10))
    style.configure("TButton", font=("Helvetica", 10))
    
    ttk.Label(app, text="CNPJ").grid(row=0, column=0, padx=10, pady=5, sticky='e')
    entry_cnpj = ttk.Entry(app)
    entry_cnpj.grid(row=0, column=1, padx=10, pady=5, sticky='w')
    ttk.Button(app, text="Buscar", command=buscar_razao_social).grid(row=0, column=2, padx=10, pady=5, sticky='w')
    
    ttk.Label(app, text="Razão Social").grid(row=1, column=0, padx=10, pady=5, sticky='e')
    entry_razao_social = ttk.Entry(app)
    entry_razao_social.grid(row=1, column=1, padx=10, pady=5, sticky='w')
    
    ttk.Label(app, text="Administrador").grid(row=2, column=0, padx=10, pady=5, sticky='e')
    entry_administrador = ttk.Entry(app)
    entry_administrador.grid(row=2, column=1, padx=10, pady=5, sticky='w')
    
    ttk.Label(app, text="Identidade").grid(row=3, column=0, padx=10, pady=5, sticky='e')
    entry_identidade = ttk.Entry(app)
    entry_identidade.grid(row=3, column=1, padx=10, pady=5, sticky='w')
    
    ttk.Label(app, text="CPF").grid(row=4, column=0, padx=10, pady=5, sticky='e')
    entry_cpf = ttk.Entry(app)
    entry_cpf.grid(row=4, column=1, padx=10, pady=5, sticky='w')
    
    ttk.Label(app, text="Módulos").grid(row=5, column=0, padx=10, pady=5, sticky='e')
    modulos_frame = tk.Frame(app)
    modulos_frame.grid(row=5, column=1, padx=10, pady=5, sticky='w')
    
    modulos_var = {
        "Fiscal/Contábil": tk.BooleanVar(),
        "Pessoal": tk.BooleanVar(),
        "Módulo X": tk.BooleanVar(),
        "Módulo Y": tk.BooleanVar(),
    }
    
    for i, (modulo, var) in enumerate(modulos_var.items()):
        tk.Checkbutton(modulos_frame, text=modulo, variable=var).grid(row=0, column=i, padx=5, pady=5, sticky='w')
    
    ttk.Label(app, text="Valor da Parcela").grid(row=6, column=0, padx=10, pady=5, sticky='e')
    entry_valor_parcela = ttk.Entry(app)
    entry_valor_parcela.grid(row=6, column=1, padx=10, pady=5, sticky='w')
    
    ttk.Label(app, text="Número de Parcelas").grid(row=7, column=0, padx=10, pady=5, sticky='e')
    entry_numero_parcelas = ttk.Entry(app)
    entry_numero_parcelas.grid(row=7, column=1, padx=10, pady=5, sticky='w')

    ttk.Label(app, text="Email").grid(row=8, column=0, padx=10, pady=5, sticky='e')
    entry_email = ttk.Entry(app)
    entry_email.grid(row=8, column=1, padx=10, pady=5, sticky='w')
    
    ttk.Button(app, text="Gerar Documento", command=lambda: gerar_documento(
        entry_cnpj.get(), entry_razao_social.get(), entry_administrador.get(),
        entry_identidade.get(), entry_cpf.get(),
        ", ".join([modulo for modulo, var in modulos_var.items() if var.get()]),
        entry_valor_parcela.get(), entry_numero_parcelas.get(), entry_email.get()
    )).grid(row=9, column=0, columnspan=3, padx=10, pady=20)

    app.mainloop()

root = tk.Tk()
root.withdraw()
icon=PhotoImage(file=r"/path-yo-your-image-icon.png") # Edit here
root.iconphoto(True, icon)
login()
root.mainloop()
